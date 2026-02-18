# =====================================================================
# TOOLS - External integrations for PO generation, email, and Slack
# =====================================================================
# Each tool extends BaseTool and returns a ToolResult (defined in 02_core).
# Tools are "dumb" executors: they do not decide *what* to do, only *how*.
# The decision logic lives in the agents (04_agents.py).
#
# Tools provided:
#   POGenerator - Fills .docx templates with PO data fields
#   EmailTool   - Reads unread emails (IMAP) and sends emails (SMTP)
#   SlackTool   - Posts messages, uploads files, and polls for new
#                 messages from a Slack channel
# =====================================================================
import smtplib
import imaplib
import email
from email.message import EmailMessage
from docx import Document
# Slack SDK is optional.  If not installed, SlackTool falls back to mock
# mode (prints messages to stdout instead of posting to a real channel).
try:
    from slack_sdk import WebClient
    from slack_sdk.errors import SlackApiError
    HAS_SLACK = True
except ImportError:
    HAS_SLACK = False
    print("slack_sdk not found. SlackTool will default to mock mode.")

# Base class that all tools must implement.
class BaseTool(ABC):
    name: str = "base_tool"
    description: str = "Base tool description"

    @abstractmethod
    def execute(self, **kwargs) -> ToolResult:
        """Execute the tool's action."""
        pass

# 1. PO Generator
# =================
# Opens a .docx template, replaces placeholder keys (e.g., X_SellerName)
# with actual data values, and saves the result to the outputs/ directory.
# Searches recursively if the template is not in the root Templates folder.
class POGenerator(BaseTool):
    name: str = "po_generator"
    description: str = "Generates Purchase Orders or Packing Lists by filling .docx templates."

    def __init__(self, templates_dir: str = "Templates"):
        self.templates_dir = templates_dir
        self.output_dir = "outputs"
        os.makedirs(self.output_dir, exist_ok=True)

    def execute(self, template_name: str, data: dict) -> ToolResult:
        template_path = os.path.join(self.templates_dir, template_name)
        if not os.path.exists(template_path):
             found = False
             for root, dirs, files in os.walk(self.templates_dir):
                 if template_name in files:
                     template_path = os.path.join(root, template_name)
                     found = True
                     break
             if not found:
                 return ToolResult(success=False, output=None, error=f"Template '{template_name}' not found in {self.templates_dir}")

        try:
            doc = Document(template_path)
            # Basic paragraph replace
            for paragraph in doc.paragraphs:
                for key, value in data.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, str(value))
            # Basic table replace
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in data.items():
                                if key in paragraph.text:
                                    paragraph.text = paragraph.text.replace(key, str(value))

            filename = f"Generated_{template_name}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            output_path = os.path.join(self.output_dir, filename)
            doc.save(output_path)
            
            return ToolResult(success=True, output={
                "message": "Document generated successfully",
                "file_path": os.path.abspath(output_path),
                "file_name": filename
            })
        except Exception as e:
            return ToolResult(success=False, output=None, error=str(e))

# 2. Communication Tools
# ========================
# EmailTool wraps IMAP (for reading unread emails) and SMTP (for sending).
# If credentials are missing, it operates in mock mode and returns dummy
# data so the system can still be tested without a real mailbox.
class EmailTool(BaseTool):
    name: str = "email_tool"
    description: str = "Reads unread emails and sends emails using SMTP/IMAP."

    def __init__(self, email_address: str, password: str, smtp_server="smtp.gmail.com", imap_server="imap.gmail.com"):
        self.email_address = email_address
        self.password = password
        self.smtp_server = smtp_server
        self.imap_server = imap_server

    def execute(self, action: str = "fetch_unread", **kwargs) -> ToolResult:
        if action == "fetch_unread":
            return self.fetch_unread()
        elif action == "send_email":
            return self.send_email(kwargs.get("to_email"), kwargs.get("subject"), kwargs.get("body"))
        else:
            return ToolResult(success=False, error=f"Unknown action: {action}")

    def fetch_unread(self) -> ToolResult:
        if not self.email_address or not self.password:
             return ToolResult(success=True, output=[{"subject": "Mock Lead", "sender": "test@user.com", "body": "I need 500kg of Coffee."}])

        try:
            mail = imaplib.IMAP4_SSL(self.imap_server)
            mail.login(self.email_address, self.password)
            mail.select('inbox')
            status, messages = mail.search(None, 'UNSEEN')
            email_ids = messages[0].split()
            logging.info(f"Connected to {self.email_address}. Found {len(email_ids)} unread emails.")
            
            results = []
            for e_id in email_ids[-3:]:
                _, msg_data = mail.fetch(e_id, '(RFC822)')
                for response_part in msg_data:
                    if isinstance(response_part, tuple):
                        msg = email.message_from_bytes(response_part[1])
                        subject = msg['subject']
                        sender = msg['from']
                        message_id = msg['Message-ID']
                        body = ""
                        if msg.is_multipart():
                            for part in msg.walk():
                                if part.get_content_type() == "text/plain":
                                    body = part.get_payload(decode=True).decode()
                                    break
                        else:
                            body = msg.get_payload(decode=True).decode()
                        results.append({"subject": subject, "sender": sender, "body": body, "message_id": message_id})
            
            mail.close()
            mail.logout()
            return ToolResult(success=True, output=results)
        except Exception as e:
            logging.error(f"Email fetch failed: {e}")
            return ToolResult(success=False, output=[], error=str(e))

    def send_email(self, to_email: str, subject: str, body: str, attachment_path: str = None) -> ToolResult:
        """Send an email with optional HTML formatting and file attachment.

        Args:
            to_email: Recipient email address.
            subject: Email subject line.
            body: Plain-text body (newlines are converted to <br> for HTML).
            attachment_path: Optional absolute path to a file to attach.
        """
        if not self.email_address or not self.password:
             print(f"[MOCK EMAIL] To: {to_email} | Subj: {subject}")
             if attachment_path:
                 print(f"[MOCK EMAIL] Attachment: {attachment_path}")
             return ToolResult(success=True, output="Mock email sent")
        try:
            from email.mime.multipart import MIMEMultipart
            from email.mime.text import MIMEText
            from email.mime.base import MIMEBase
            from email import encoders

            msg = MIMEMultipart()
            msg['Subject'] = subject
            msg['From'] = self.email_address
            msg['To'] = to_email

            # Convert plain-text body to simple HTML for clean rendering
            html_body = body.replace("\n", "<br>\n")
            html_content = f"""<html><body style="font-family: Arial, sans-serif; font-size: 14px; line-height: 1.6; color: #333;">
{html_body}
</body></html>"""
            msg.attach(MIMEText(html_content, 'html'))

            # Attach file if provided (e.g., the generated PO document)
            if attachment_path and os.path.exists(attachment_path):
                with open(attachment_path, 'rb') as f:
                    part = MIMEBase('application', 'octet-stream')
                    part.set_payload(f.read())
                encoders.encode_base64(part)
                filename = os.path.basename(attachment_path)
                part.add_header('Content-Disposition', f'attachment; filename="{filename}"')
                msg.attach(part)

            with smtplib.SMTP_SSL(self.smtp_server, 465) as smtp:
                smtp.login(self.email_address, self.password)
                smtp.send_message(msg)
            return ToolResult(success=True, output="Email sent successfully")
        except Exception as e:
            return ToolResult(success=False, output=None, error=str(e))

# SlackTool wraps the Slack Web API for posting messages, uploading
# files, and polling for new messages.  It resolves channel names to
# IDs automatically and caches known private channel IDs.
class SlackTool(BaseTool):
    name: str = "slack_tool"
    description: str = "Posts messages and files to Slack."

    def __init__(self, token: str, default_channel: str = "#general"):
        self.token = token
        self.default_channel = default_channel
        self.client = WebClient(token=token) if HAS_SLACK and token else None

    def _resolve_channel_id(self, channel_name: str) -> str:
        if not self.client: return channel_name
        if not channel_name.startswith("#"): return channel_name
        KNOWN_PRIVATE_CHANNELS = { "#agenticai-group-9": "C0AC5C7UMST" }
        if channel_name in KNOWN_PRIVATE_CHANNELS: return KNOWN_PRIVATE_CHANNELS[channel_name]

        try:
            clean_name = channel_name.lstrip("#")
            response = self.client.conversations_list(types="public_channel,private_channel", limit=1000)
            for channel in response.get("channels", []):
                if channel["name"] == clean_name: return channel["id"]
            return channel_name 
        except SlackApiError as e:
            print(f"ERROR: Failed to resolve channel ID: {e}")
            return channel_name

    def execute(self, action: str, **kwargs) -> ToolResult:
        if action == "post_message":
            return self.post_message(kwargs.get("text"), kwargs.get("channel"), kwargs.get("blocks"))
        elif action == "upload_file":
            return self.upload_file(kwargs.get("filepath"), kwargs.get("comment"), kwargs.get("channel"))
        else:
            return ToolResult(success=False, error=f"Unknown action: {action}")

    def post_message(self, text: str, channel: str = None, blocks: List = None) -> ToolResult:
        if not self.client:
            print(f"[MOCK SLACK] {text}")
            return ToolResult(success=True, output="Mock message posted")
        try:
            target_name = channel or self.default_channel
            target_id = self._resolve_channel_id(target_name)
            response = self.client.chat_postMessage(channel=target_id, text=text, blocks=blocks)
            return ToolResult(success=True, output=response.data)
        except SlackApiError as e:
            return ToolResult(success=False, output=None, error=str(e))

    def upload_file(self, filepath: str, comment: str, channel: str = None) -> ToolResult:
        if not self.client:
             print(f"[MOCK SLACK FILE] Uploading {filepath}: {comment}")
             return ToolResult(success=True, output="Mock file uploaded")
        try:
            target_name = channel or self.default_channel
            target_id = self._resolve_channel_id(target_name)
            response = self.client.files_upload_v2(channel=target_id, file=filepath, initial_comment=comment)
            return ToolResult(success=True, output=response.data)
        except SlackApiError as e:
             return ToolResult(success=False, output=None, error=str(e))

    def fetch_messages(self, channel: str, oldest: str = None) -> ToolResult:
        if not self.client: return ToolResult(success=True, output=[])
        try:
            target_name = channel or self.default_channel
            target_id = self._resolve_channel_id(target_name)
            response = self.client.conversations_history(channel=target_id, limit=10)
            all_messages = response.data.get("messages", [])
            
            if oldest:
                oldest_float = float(oldest)
                messages = [m for m in all_messages if float(m.get('ts', 0)) > oldest_float]
            else:
                messages = all_messages
            return ToolResult(success=True, output=messages)
        except SlackApiError as e:
            if e.response.get("error") == "missing_scope":
                print(f"[SLACK PERMISSION ERROR] Missing Scopes.")
            return ToolResult(success=False, output=[], error=str(e))
